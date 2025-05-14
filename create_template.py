from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name, size, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def create_template():
    # Create a new Document
    doc = Document()
    
    # Define colors
    heading_color = (33, 33, 33)  # Very dark gray, almost black
    
    # Add title
    title = doc.add_heading('Konsultprofil', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    # Add template variables
    name_para = doc.add_paragraph()
    name_run = name_para.add_run('{{name}}')
    set_font(name_run, 'Roboto', 26, True, heading_color)
    
    slogan_para = doc.add_paragraph()
    slogan_run = slogan_para.add_run('{{slogan}}')
    set_font(slogan_run, 'Roboto', 20, True, heading_color)
    
    doc.add_paragraph()
    
    # Summary section
    summary_heading = doc.add_heading('Sammanfattning', level=1)
    for run in summary_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run('{{summary}}')
    set_font(summary_run, 'Roboto', 11)
    doc.add_paragraph()
    
    # Expertise section
    expertise_heading = doc.add_heading('Expertis', level=1)
    for run in expertise_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    expertise_para = doc.add_paragraph()
    expertise_run = expertise_para.add_run('{% for skill in expertise %}{{skill}}{% if not loop.last %}, {% endif %}{% endfor %}')
    set_font(expertise_run, 'Roboto', 11)
    doc.add_paragraph()
    
    # Assignments section
    assignments_heading = doc.add_heading('Uppdrag', level=1)
    for run in assignments_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    doc.add_paragraph('{% for assignment in assignment %}')
    
    # Assignment details
    for field in ['Roll', 'Kund', 'Period', 'Beskrivning', 'Approach']:
        field_heading = doc.add_heading(field, level=2)
        for run in field_heading.runs:
            set_font(run, 'Roboto', 20, True, heading_color)
        field_para = doc.add_paragraph('{{assignment.' + field.lower() + '}}')
        for run in field_para.runs:
            set_font(run, 'Roboto', 11)
    
    doc.add_paragraph('{% endfor %}')
    
    # Contact information
    contact_heading = doc.add_heading('Kontaktinformation', level=1)
    for run in contact_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    for field in ['Email', 'Telefon', 'LinkedIn', 'Plats']:
        field_heading = doc.add_heading(field, level=2)
        for run in field_heading.runs:
            set_font(run, 'Roboto', 20, True, heading_color)
        field_para = doc.add_paragraph('{{' + field.lower() + '}}')
        for run in field_para.runs:
            set_font(run, 'Roboto', 11)
    
    doc.add_paragraph()
    
    # Technical skills
    tech_heading = doc.add_heading('Teknisk kompetens', level=1)
    for run in tech_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    for field in ['Teknologier', 'Metoder', 'Verktyg']:
        field_heading = doc.add_heading(field, level=2)
        for run in field_heading.runs:
            set_font(run, 'Roboto', 20, True, heading_color)
        field_para = doc.add_paragraph('{% for item in ' + field.lower() + ' %}{{item}}{% if not loop.last %}, {% endif %}{% endfor %}')
        for run in field_para.runs:
            set_font(run, 'Roboto', 11)
    
    doc.add_paragraph()
    
    # Languages
    lang_heading = doc.add_heading('Språk', level=1)
    for run in lang_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    lang_para = doc.add_paragraph('{% for lang in languages %}{{lang.language}}: {{lang.level}}{% if not loop.last %}\n{% endif %}{% endfor %}')
    for run in lang_para.runs:
        set_font(run, 'Roboto', 11)
    doc.add_paragraph()
    
    # Education
    edu_heading = doc.add_heading('Utbildning', level=1)
    for run in edu_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    doc.add_paragraph('{% for edu in education %}')
    for field in ['Institution', 'Fokus', 'Period']:
        field_heading = doc.add_heading(field, level=2)
        for run in field_heading.runs:
            set_font(run, 'Roboto', 20, True, heading_color)
        field_para = doc.add_paragraph('{{edu.' + field.lower() + '}}')
        for run in field_para.runs:
            set_font(run, 'Roboto', 11)
    doc.add_paragraph('{% endfor %}')
    
    # Certifications
    cert_heading = doc.add_heading('Certifieringar', level=1)
    for run in cert_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    cert_para = doc.add_paragraph('{% for cert in certifications %}{{cert}}{% if not loop.last %}, {% endif %}{% endfor %}')
    for run in cert_para.runs:
        set_font(run, 'Roboto', 11)
    doc.add_paragraph()
    
    # Employment information
    emp_heading = doc.add_heading('Anställningsinformation', level=1)
    for run in emp_heading.runs:
        set_font(run, 'Roboto', 56, True, heading_color)
    
    for field in ['Anställningstyp', 'Anställd av']:
        field_heading = doc.add_heading(field, level=2)
        for run in field_heading.runs:
            set_font(run, 'Roboto', 20, True, heading_color)
        field_para = doc.add_paragraph('{{' + field.lower().replace(' ', '_') + '}}')
        for run in field_para.runs:
            set_font(run, 'Roboto', 11)
    
    # Save the template
    doc.save('konsultprofil_template.docx')

if __name__ == '__main__':
    create_template() 